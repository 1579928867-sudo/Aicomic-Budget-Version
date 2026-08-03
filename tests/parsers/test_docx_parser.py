"""Tests for DocxParser — Word document text extraction."""

from pathlib import Path
import pytest
from aicomic.parsers.docx_parser import DocxParser


class TestDocxParserSupports:
    """supports() checks file extension only."""

    def test_supports_docx(self):
        assert DocxParser.supports(Path("chapter.docx")) is True

    def test_supports_uppercase(self):
        assert DocxParser.supports(Path("CHAPTER.DOCX")) is True

    def test_rejects_txt(self):
        assert DocxParser.supports(Path("chapter.txt")) is False

    def test_rejects_pdf(self):
        assert DocxParser.supports(Path("chapter.pdf")) is False


class TestDocxParserParse:
    """parse() extracts paragraph text from .docx files."""

    def test_single_paragraph(self, tmp_path):
        from docx import Document
        file_path = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("萧澈缓缓睁开眼睛。")
        doc.save(str(file_path))

        parser = DocxParser()
        result = parser.parse(file_path)
        assert "萧澈缓缓睁开眼睛。" in result

    def test_multiple_paragraphs(self, tmp_path):
        from docx import Document
        file_path = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("第一段：清晨暖光透过窗棂。")
        doc.add_paragraph("第二段：红色帷帐在风中飘动。")
        doc.add_paragraph("第三段：远处隐约传来鸟鸣。")
        doc.save(str(file_path))

        parser = DocxParser()
        result = parser.parse(file_path)
        # Paragraphs should be separated by double newlines
        assert "第一段：清晨暖光透过窗棂。" in result
        assert "第二段：红色帷帐在风中飘动。" in result
        assert "第三段：远处隐约传来鸟鸣。" in result
        # Verify paragraph separation
        assert "\n\n" in result

    def test_empty_document(self, tmp_path):
        from docx import Document
        file_path = tmp_path / "empty.docx"
        doc = Document()
        doc.save(str(file_path))

        parser = DocxParser()
        result = parser.parse(file_path)
        assert result == ""

    def test_skips_empty_paragraphs(self, tmp_path):
        from docx import Document
        file_path = tmp_path / "sparse.docx"
        doc = Document()
        doc.add_paragraph("第一段内容。")
        doc.add_paragraph("")  # empty paragraph
        doc.add_paragraph("")  # another empty
        doc.add_paragraph("第三段内容。")
        doc.save(str(file_path))

        parser = DocxParser()
        result = parser.parse(file_path)
        # Empty paragraphs should be filtered out
        assert result.count("\n\n") == 1  # exactly one separation
        assert "第一段内容。" in result
        assert "第三段内容。" in result

    def test_corrupted_file_raises(self, tmp_path):
        """A file with .docx extension that isn't a valid docx should error."""
        file_path = tmp_path / "fake.docx"
        file_path.write_text("not a real docx file", encoding="utf-8")
        parser = DocxParser()
        with pytest.raises(ValueError, match="Failed to parse"):
            parser.parse(file_path)
